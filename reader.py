"""
Turning a file into clean text, page by page.

Two things matter here and are easy to get wrong:

  1. Page numbers are captured while we are already walking the
     document. They cannot be recovered afterwards without
     re-processing everything, and they are what lets a citation
     name a page rather than only a file.

  2. Word documents keep paragraphs and tables in separate places.
     Code that reads only paragraphs gets nothing from a document
     whose content sits in tables — no error, just an empty result.
"""

import hashlib
import re
from collections import Counter
from pathlib import Path

import pdfplumber
import docx


def file_hash(path):
    """Fingerprint of the file contents, so the same file is never added twice."""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for block in iter(lambda: f.read(65536), b""):
            h.update(block)
    return h.hexdigest()


# ─── PDF ─────────────────────────────────────────────────────────

def read_pdf(path):
    """Returns [(page_number, text), ...] — one entry per page."""
    pages = []
    with pdfplumber.open(path) as pdf:
        for n, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if not text.strip():
                text = _read_two_columns(page)
            pages.append((n, text))
    return pages


def _read_two_columns(page):
    """
    Fallback for pages where reading straight across produces nothing
    useful. Splits the page down the middle and reads each half in turn,
    so sentences in a two-column layout do not interleave.
    """
    try:
        mid = page.width / 2
        left = page.crop((0, 0, mid, page.height)).extract_text() or ""
        right = page.crop((mid, 0, page.width, page.height)).extract_text() or ""
        return (left + "\n" + right).strip()
    except Exception:
        return ""


# ─── Word ────────────────────────────────────────────────────────

def read_docx(path):
    """
    Word files have no pages until they are rendered, so everything
    is recorded as page 1. Table content is collected as well as
    paragraphs — some documents keep all their text in tables.
    """
    d = docx.Document(path)

    parts = [p.text for p in d.paragraphs if p.text.strip()]

    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return [(1, "\n".join(parts))]


# ─── cleaning ────────────────────────────────────────────────────

def strip_repeats(pages, threshold=0.5):
    """
    Remove lines that appear on most pages — journal names across the
    top, page numbers at the foot. Detected by frequency rather than
    position, because position varies between documents.
    """
    if len(pages) < 3:
        return pages

    seen = Counter()
    for _, text in pages:
        for line in set(text.split("\n")):
            line = line.strip()
            if 0 < len(line) < 80:
                seen[line] += 1

    cutoff = len(pages) * threshold
    junk = {line for line, n in seen.items() if n > cutoff}

    cleaned = []
    for page_no, text in pages:
        kept = [ln for ln in text.split("\n") if ln.strip() not in junk]
        cleaned.append((page_no, "\n".join(kept)))
    return cleaned


def tidy(text):
    """Collapse runaway whitespace and join hyphenated line breaks."""
    text = re.sub(r"-\n(\w)", r"\1", text)      # word split across a line
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ─── the one function the pipeline calls ─────────────────────────

def read(path):
    """
    Returns (pages, note).

    pages is [(page_number, text), ...]
    note is None if all went well, or a reason the file could not be used.
    """
    path = Path(path)
    suffix = path.suffix.lower()

    try:
        if suffix == ".pdf":
            pages = read_pdf(path)
        elif suffix in (".docx", ".doc"):
            pages = read_docx(path)
        else:
            return [], f"unsupported file type: {suffix}"
    except Exception as e:
        return [], f"could not open: {type(e).__name__}"

    if not pages:
        return [], "no pages found"

    pages = strip_repeats(pages)
    pages = [(n, tidy(t)) for n, t in pages]
    return pages, None
